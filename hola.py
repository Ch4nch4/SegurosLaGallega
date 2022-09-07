import docx
from docx.shared import Inches,Cm,RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT,WD_BREAK


documento = docx.Document()
for seccion in documento.sections:
    seccion.left_margin = Cm(2)
    seccion.right_margin = Cm(2)
seccion = documento.sections[0]

seccion_titulo = seccion.header

encabezado = seccion_titulo.add_paragraph()
encabezado_run = encabezado.add_run()
encabezado_run.add_text("\t\t\t\t\t\tRegistro de alumno")
encabezado_run.add_text("\n__________________________________________________________________________________________________________________________")

t = "Nombre:"
t1 = "MiNombre"     
tabla = documento.add_table(rows=4,cols=2)
celda = tabla.rows[0].cells

tabla_celda = celda[0].add_paragraph().add_run()
p= tabla_celda.add_text(t)

p1 = tabla_celda.add_text(t1)


documento.save("Nuevo.docx")