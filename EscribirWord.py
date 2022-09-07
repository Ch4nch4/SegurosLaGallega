#escribir en word
from pydoc import Doc
import docx
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT,WD_BREAK

doc= docx.Document()
doc.add_paragraph('Listado general de personas')
#Agregando encabezados
doc.add_heading('Header 0', 0)
doc.add_heading('Header 1', 1)
doc.add_heading('Header 2', 2)
doc.add_heading('Header 3', 3)
doc.add_heading('Header 4', 4)

#agregando saltos de linea y de pagina
doc.add_paragraph('This is on the first page!')
doc.add_heading('Imagen de la Luna', 0)
doc.add_picture('luna.jpg', 
                width=docx.shared.Inches(1),
                height=docx.shared.Cm(4))

#salto de pagina
parrafo = doc.add_paragraph("Primera línea")
parrafo.add_run().add_break()                #WD_BREAK.LINE, es el valor por defecto.
parrafo.add_run("Segunda línea")
parrafo.add_run().add_break(WD_BREAK.PAGE)    #salto pagina

doc.add_paragraph('This is on the second page!')


#guardo el archivo word
doc.save('listadoPersonas.docx')


