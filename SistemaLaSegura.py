import sys
from PyQt5 import QtWidgets, uic
from CapaNegocio.Autos import *
from CapaNegocio.Persona import *
from CapaNegocio.Siniestros import *
from CapaNegocio.Vincular import *
from datetime import date,datetime
import openpyxl
import re
from pydoc import Doc
import docx
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT,WD_BREAK
import os

class Iniciar:
    def __init__(self):
        app=QtWidgets.QApplication([])
        self.ventana = uic.loadUi(".\\ClaseU1\\SegurosLaGallega\\CapaVista\\ventanaPpal.ui")
        self.ventana.show()
 
        ####### conectamos menu con funcion
        self.ventana.actionSalir.triggered.connect(self.accionSalir)
        self.ventana.actionAutos.triggered.connect(self.accionCRUDAutos)
        self.ventana.actionSiniestros.triggered.connect(self.accionSiniestros)
        
        app.exec()
        
# Manejo de Menus de opciones 
    def accionSalir(self):
        sys.exit()
    
    def accionCRUDAutos(self):
        self.crudAutos = uic.loadUi(".\\ClaseU1\\SegurosLaGallega\\CapaVista\\crudAutos.ui")
        self.crudAutos.show()

        ###### conectamos los botones 
        self.crudAutos.btnAgregar.clicked.connect(self.Click_btnAgregar)
        self.crudAutos.btnBorrar.clicked.connect(self.Click_btnBorrar)
        self.crudAutos.btnBuscar.clicked.connect(self.Click_btnBuscar)  
        self.crudAutos.btnActualizar.clicked.connect(self.Click_btnActualizar) 
        
    def accionSiniestros(self):
        self.ventanaSiniestros = uic.loadUi(".\\ClaseU1\\SegurosLaGallega\\CapaVista\\ventanaAseguradora.ui")
        self.ventanaSiniestros.show()
        ###### conectamos los botones 
        self.ventanaSiniestros.btnBuscarPersona.clicked.connect(self.Click_btnBuscarPersona)
        self.ventanaSiniestros.btnBuscarVehiculo.clicked.connect(self.Click_btnBuscarVehiculo)
        self.ventanaSiniestros.btnAgregarInforme.clicked.connect(self.Click_btnAgregarInforme) 
        self.ventanaSiniestros.btnExcel.clicked.connect(self.Click_btnExcel)
        self.ventanaSiniestros.btnPersonasaWord.clicked.connect(self.Clicl_btnPersonasaWord)
       # self.llenar_combos()
        
            
    def llenar_combos(self):
        self.ventanaSiniestros.comboBoxVehiculo.clear()
        self.ventanaSiniestros.comboBoxCarnet.clear()
        tautos=Autos()
        autos=tautos.getTodos()
        for auto in autos:
            self.ventanaSiniestros.comboBoxVehiculo.addItems(auto[0][0])
            self.ventanaSiniestros.listView.addItems(str(auto[0]))
            
        tpersonas=Personas()
        personas=tpersonas.getTodos()
        self.ventanaSiniestros.tableview.columns('carnet','nombe','dir')
        for persona in personas:
            self.ventanaSiniestros.comboBoxCarnet.addItems(str(persona[1]))
            self.self.ventanaSiniestros.tableview.addItems(str(persona[1]),str(persona[2]),str(persona[4]))
        
            
 
# Manejo de Botones del accionCRUDAutos
    def Click_btnAgregar(self):
        # Si en el texto de la matricula no hay nada INFORMAR un error
        ########## verificar si la matricula tiene la forma AAA111 o AA111AA
        # Si en el texto de la modelo no hay nada INFORMAR un error
        # Si en el texto de la año no hay nada INFORMAR un error
        ###### verificar que el año sea <= añoactual
        # llamar a la capa negocio para insertar el auto

        ######### obtener los valores del los input 
        matricula = self.crudAutos.txtMatricula.text()
        modelo = self.crudAutos.txtModelo.text()
        anio = int(self.crudAutos.txtAnio.text())

        if self.verificarMatricula(matricula) and self.verificarModelo(modelo) and self.verificarAnio(anio):
            tAutos = Autos()
            tAutos.insertAuto(matricula,modelo,anio)
            self.crudAutos.lblDisplay.setText("El auto fue agregado exitosamente")
        else:
            self.crudAutos.lblDisplay.setText("Revise los datos ingresados")
         
    def verificarMatricula(self,matricula):
        patronPatente = "[A-Z]{3}\d{3}|[A-Z]{2}\d{3}[A-Z]{2}"
        if (matricula == ""):
            self.crudAutos.lblDisplay.setText("Ingrese una matricula")
            return False
        else:
            objetoMatch=re.match(patronPatente,matricula)
            if ( objetoMatch == None ):
                self.crudAutos.lblDisplay.setText("El formato no concuerda con una patente argentina")
                return False
            else:
                return True

    def verificarModelo(self,modelo):
        if modelo == "":
           self.crudAutos.lblDisplay.setText("Ingrese un modelo")
           return False
        else:
            return True

    def verificarAnio(self,anio):
        if anio == "":
           self.crudAutos.lblDisplay.setText("Ingrese un AÑO")
           return False
        elif (anio <= date.today().year):
            return True
        else:
            return False
    
    def Click_btnBorrar(self):
        # Si en el texto de la matricula no hay nada INFORMAR un error
        # Si hay algo pero no existe en la table  INFORMAR no exiist
        #             si si existe borrarla
        matricula=self.crudAutos.txtMatricula.text()
        if self.verificarMatricula(matricula):
            if self.existeMatricula(matricula):
                tauto=Autos()
                tauto.deleteAuto(matricula)
                self.crudAutos.lblDisplay.setText("El auto fue borrado exitosamente")
            else:
                self.crudAutos.lblDisplay.setText("La matricula ingresada no existe en la Base de Datos")
        else:    
            self.crudAutos.lblDisplay.setText("La matricula no se verifica")

    def existeMatricula(self,matricula):
        tauto=Autos()
        auto= tauto.getAuto(matricula)
        if auto:
           return True
        else:
           return False

    def Click_btnBuscar(self):
        # Si en el texto de la matricula no hay nada INFORMAR un error
        # Si hay algo pero no existe en la table  INFORMAR no exiiste
        #             si si existe llenar el resto de los lineEdit
        matricula=self.crudAutos.txtMatricula.text()
        if self.verificarMatricula(matricula):
            if self.existeMatricula(matricula):
                tauto=Autos()
                auto = tauto.getAuto(matricula)
                self.crudAutos.txtModelo.setText(auto[1])
                self.crudAutos.txtAnio.setText(auto[2])
                self.crudAutos.lblDisplay.setText("El auto fue encontrado exitosamente")
            else:
                self.crudAutos.lblDisplay.setText("La matricula ingresada no existe en la Base de Datos")
        else:    
            self.crudAutos.lblDisplay.setText("La matricula no se verifica")
       
    def Click_btnActualizar(self):
        # Si en el texto de la matricula no hay nada INFORMAR un error
        # ##### la matricula tiene que existir para poder actualizarla
        # Si en el texto de la modelo no hay nada INFORMAR un error
        # Si en el texto de la año no hay nada INFORMAR un error
        # llamar a la capa negocio para actualizar el auto

        ######### obtener los valores del los input 
        matricula = self.crudAutos.txtMatricula.text()
        modelo = self.crudAutos.txtModelo.text()
        anio = int(self.crudAutos.txtAnio.text())

        if self.verificarMatricula(matricula) and self.verificarModelo(modelo) and self.verificarAnio(anio):
            if self.existeMatricula(matricula):
                tAutos = Autos()
                tAutos.updateAuto(matricula,modelo,anio)
                self.crudAutos.lblDisplay.setText("El auto fue actualizado exitosamente")
        else:
            self.crudAutos.lblDisplay.setText("Revise los datos ingresados") 

    def Click_btnBuscarPersona(self):
        carnet= self.ventanaSiniestros.txtCarnetPersona.text()
        tpersonas = Personas()
        persona=tpersonas.getPersonaxCarnet(carnet)
        if persona != None:
            self.ventanaSiniestros.txtNombreyApellidoPersona.setText(persona[2] + ' ' + persona[3])
            self.ventanaSiniestros.txtDireccionpersona.setText(persona[4])
            self.ventanaSiniestros.lblIDPERSONA.setText(str(persona[0]))
        
        # si la esa persona tiene auto ya cargo la matricula
        idpers=int(self.ventanaSiniestros.lblIDPERSONA.text())
        
        listaMatriculas=tpersonas.getMatriculasDeAutosDePersona(idpers)
        if len(listaMatriculas) > 0:
            if len(listaMatriculas) == 1:
                self.ventanaSiniestros.txtMatricula.setText(str(listaMatriculas[0][0]))
                self.ventanaSiniestros.lblautosdepersona.setText(" la persona tiene un auto de matricula " + str(listaMatriculas[0][0]) )
            else:
                strMatriculas=""
                for matricula in listaMatriculas:
                    strMatriculas = strMatriculas + ' ' + matricula
                self.ventanaSiniestros.lblautosdepersona.setText(" la persona tiene estos autos " + strMatriculas )
        # tambien establezco el nroi¿nforme que sigue
        tsinistro=Siniestro()
        nroInforme = self.ventanaSiniestros.txtnroInforme.setText(str(tsinistro.traerProximoNroInforme()))
        self.ventanaSiniestros.txtFechaInforme.setText(str(date.today()))
            
    def Click_btnBuscarVehiculo(self):
        matricula = self.ventanaSiniestros.txtMatricula.text()
        tautos = Autos()
        auto=tautos.getAuto(matricula)
        if auto != None:
            self.ventanaSiniestros.txtModelo.setText(auto[1])
            self.ventanaSiniestros.txtAnio.setText(auto[2])      
            
    def Click_btnAgregarInforme(self):
        matricula = self.ventanaSiniestros.txtMatricula.text()
        idPersona = self.ventanaSiniestros.lblIDPERSONA.text()
        nroInforme = self.ventanaSiniestros.txtnroInforme.text()
        fecha= self.ventanaSiniestros.txtFechaInforme.text()
        lugar = self.ventanaSiniestros.txtLugarInforme.text()
        importe = float(self.ventanaSiniestros.txtImporteDanios.text())
        tsiniestro = Siniestro()
        tsiniestro.insertarInforme(nroInforme,fecha,lugar)
        # si no existe la relacion vincular  de la persona con el auto vincular
        tvinculo=Vincular()
        if not tvinculo.existeVinculoPersonAuto(idPersona,matricula):
           tvinculo.vincularPersonaAuto(idPersona,matricula)
        #vincular informe persona y auto    
        tvinculo.vincularAutoInforme(matricula,nroInforme,importe)  
    
    def Click_btnExcel(self) :
        libro=openpyxl.Workbook() 
        #fijar la hoja
        hoja=libro.active
    
        hoja['A1'] = 'Listado de Autos'
        hoja._add_row()
        
        hoja.append(['Matricula','Modelo','Anio'])
        objAuto=Autos()
        autos=objAuto.getTodos()

        for fila in autos:      
            hoja.append(fila)   #cada fila es un empleado

        libro.save('listadoAuto.xlsx')
    
    def Clicl_btnPersonasaWord(self):
        doc= docx.Document()
        
        #Agregando encabezados
        doc.add_heading('Listado general de personas', 0)
        
        tpersonas=Personas()
        personas= tpersonas.getTodos()
        doc.add_paragraph('CARNET       NOMMBRE                             DIRECCION')
        for persona in personas:
            doc.add_paragraph('{}       {}, {}        {}'.format(persona[1],persona[3],persona[2],persona[4]))
       
        #guardo el archivo word
        doc.save('listadoPersonas.docx')
        
        os.startfile('listadoPersonas.docx',"print")
        
########### main
Iniciar()


 

  
 

 
 
 


 